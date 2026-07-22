from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

import build_tri_summary as base


OUT = base.ROOT / "output" / "TRI论文进展同步_20260721.docx"


def new_page(doc):
    doc.add_page_break()


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.70)
section.bottom_margin = Inches(0.70)
section.left_margin = Inches(0.86)
section.right_margin = Inches(0.86)
section.header_distance = Inches(0.40)
section.footer_distance = Inches(0.40)
base.configure_styles(doc)
doc.styles["Normal"].font.size = Pt(9.8)
doc.styles["Normal"].paragraph_format.space_after = Pt(4)
doc.styles["Normal"].paragraph_format.line_spacing = 1.16
base.add_header_footer(section)

# Page 1: cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(82)
p.paragraph_format.space_after = Pt(8)
r = p.add_run("论文研究进展同步")
base.set_run_font(r, size=11, color=base.TEAL, bold=True, font=base.FONT_HEAD)

p = doc.add_paragraph(style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("更新世界状态不等于\n重新绑定操作目标")
base.set_run_font(r, size=27, color=base.NAVY, bold=True, font=base.FONT_HEAD)

p = doc.add_paragraph(style="Subtitle")
r = p.add_run("工具型智能体中的时序指称完整性（Temporal Referent Integrity, TRI）")
base.set_run_font(r, size=13, color=base.MUTED, font=base.FONT_HEAD)

doc.add_paragraph().paragraph_format.space_after = Pt(10)
base.add_callout(
    doc,
    "核心问题",
    "Agent 刷新外部状态时，世界事实可以更新，但已经由用户语义确定的操作对象不能在没有授权的情况下被悄悄换成新的选择器赢家。",
    base.LIGHT_TEAL,
    base.TEAL,
)

meta = doc.add_table(rows=4, cols=2)
meta.style = "Table Grid"
for row, (label, value) in zip(meta.rows, [
    ("拟投稿", "AAAI 2027 Main Technical Track"),
    ("核心方法", "刷新前目标承诺编译（Compile-then-act, CTA）"),
    ("已完成实验", "主实验、三模型复制、SQLite 写入、人类验证、外部审计"),
    ("更新时间", "2026 年 7 月 21 日"),
]):
    base.set_cell_shading(row.cells[0], base.LIGHT_BLUE)
    base.set_para_text(row.cells[0].paragraphs[0], label, size=9.5, color=base.NAVY, bold=True)
    base.set_para_text(row.cells[1].paragraphs[0], value, size=9.5)
base.set_table_geometry(meta, [1900, 7460])

new_page(doc)

# Page 2: translated abstract and introduction
doc.add_heading("1. 论文摘要（中文整理版）", level=1)
base.add_body(doc, "刷新外部状态本身并不授权工具型智能体改变先前指称所对应的实体。本文将这一绑定后问题形式化为时序指称完整性（Temporal Referent Integrity，TRI），并区分世界知识更新与指称转换授权。我们提出刷新前目标承诺编译（Compile-then-act，CTA）：在环境变化前明确记录用户是在当前状态中已经绑定具体实体，还是有意把选择推迟到刷新之后。")
base.add_body(doc, "在冻结的 160 任务诊断集上，Generic Structured Ledger 在 Qwen3.5-122B 和 GLM-5.1 上分别达到 64.4% 和 71.9%，CTA 则达到 95.0% 和 96.2%；Always-Lock 与 Always-Reevaluate 均只有 60%，且失败模式相反。在另行冻结的 240 任务复制中，给定正确初始绑定，Generic 分别在 43/72 和 38/80 个机会中发生漂移，CTA 与 Lifecycle-Gated 均为 0；DeepSeek 的后续复制得到 59/79 对 0/70。所有 Generic 核心漂移均可重放为 SQLite 错误实体写入。三名盲标者在 100 条原始或人工改写指令上取得 Fleiss kappa=.708，多数票与 gold 一致率为 86%。公开 ToolSandbox、AppWorld 与 tau3 审计几乎未发现原生严格机会，低干预工具闭环也为零结果。因此，TRI 是一种依赖模型、控制器与机会结构的受控机制诊断，而不是现实发生率或通用安全结论。")

doc.add_heading("2. 引言（中文整理版）", level=1)
base.add_body(doc, "考虑两条在同一次邮箱刷新前提出的请求：（1）“现在选出优先级最高的未读邮件；刷新邮箱后回复它。”（2）“先刷新邮箱，再选出优先级最高的未读邮件并回复。”假设邮件 A 在刷新前优先级最高，邮件 B 在刷新后成为最高。两条请求经历相同的环境变化，但第一条已将目标绑定为 A，第二条则要求在新状态中选择 B。第一条中改为回复 B 不是纠正过时信息，而是对错误实体执行操作。")
base.add_body(doc, "这种区别很容易在 LLM Agent 控制器中丢失。控制器即使保留完整历史并追加最新观察，下游模型仍可能把旧选择条件重新应用到新状态；反过来，始终保留旧 ID 又会破坏确实要求刷新后重选的请求。可靠工具使用因此需要选择性不变性：世界事实应当更新，但已经形成的目标身份只在用户授权时转换。")
base.add_body(doc, "TRI 并不等同于一般实体记忆。记住一个 ID 并不能说明它是已经承诺的指称对象，还是某个动态查询先前的输出；实体仍然存在，也不代表它对当前动作仍然有效。本文因此将 reference mode、绑定身份、selector provenance、动作有效性与 fallback 分开建模，并严格条件化到初始绑定正确，以排除首次选错、工具顺序错误和 API 失败。")
base.add_body(doc, "方法上，Generic Structured Ledger 保存身份、事实、选择器和动作前置条件，但把这些字段跨刷新如何交互留给模型推断。CTA 在刷新前先编译绑定时间和目标身份；类型化 Lifecycle 扩展进一步记录有效性与 fallback，并由确定性 gate 在执行边界保护 Preserve 分支。现有消融显示，主要收益来自刷新前承诺编译，类型化状态与 gate 主要增加可审计性和执行保护。")
base.add_body(doc, "本文的主要贡献包括：")
for item in [
    "形式化 TRI，将绑定后的目标转换授权与世界状态更新、初始实体绑定和动作有效性区分开；",
    "构建冻结的成簇诊断任务，并通过三模型复制、人工改写和 SQLite 写入审计验证现象及后果；",
    "系统比较结构化状态、完整历史、刷新后强提醒和预刷新承诺编译，同时通过公开基准审计明确外部覆盖盲区与适用边界。",
]:
    base.add_bullet(doc, item)
base.add_callout(doc, "引言主线", "状态刷新更新世界知识，但不自动授权 Agent 改变已经绑定的操作目标。", base.LIGHT_ORANGE, base.ORANGE)

# Method and experiment map
doc.add_heading("3. 方法与实验布局", level=1)
doc.add_heading("3.1 刷新前目标承诺编译", level=2)
base.add_body(doc, "普通 Structured Ledger 会保存原指令、初始 ID、实体快照、选择条件和动作前置条件，但仍让模型在刷新后自由解释这些信息。问题在于，记录了旧选择器并不等于明确了它现在只是身份来源，还是仍可执行的查询。")
base.add_numbered_list(doc, [
    "刷新前判断用户是在当前状态中已经选定对象，还是有意把选择留到刷新后；",
    "Preserve 分支保存稳定实体 ID，Reevaluate 分支保存未解析的选择条件；",
    "刷新环境只更新世界事实，不自动改写已形成的身份承诺；",
    "执行时 Preserve 使用原 ID，Reevaluate 在新状态中重新选择；目标失效时按明确策略处理。",
])
base.add_body(doc, "论文还实现了类型化 Lifecycle 状态和确定性 gate，用于显式记录 reference mode、bound target ID、selector、动作有效性和 fallback。现有消融显示，主要收益来自刷新前承诺编译，gate 主要增加可审计性和执行保护。")

doc.add_heading("3.2 已完成的实验布局", level=2)
base.add_table(doc,
    ["实验", "规模", "目的"],
    [
        ("v3 主实验", "160 条，2 模型", "验证现象、主方法和消融"),
        ("v7 独立复制", "240 条，3 模型", "新状态、新 schema、第三模型复现"),
        ("强基线", "完整历史 + 刷新后强提醒", "检查是否只是 ledger 诱发"),
        ("SQLite 后果", "直接执行 + 1,440 次重放", "验证错误是否变成真实写入"),
        ("人工验证", "3 人盲标 100 条 + 50 条改写", "验证语义和模板外稳健性"),
        ("外部验证", "ToolSandbox/AppWorld/tau3", "检查公开任务覆盖和工具接口表现"),
    ], [2100, 2350, 4910], small=8.8)
base.add_body(doc, "主实验在模型调用前冻结数据、runner、hash 和推理设置；置信区间以完整模板簇或状态簇为重采样单位，使用 10,000 次 cluster bootstrap。")

# v3 and ablations
doc.add_heading("4. 受控主实验与消融结果", level=1)
base.add_table(doc,
    ["控制器", "Qwen3.5", "GLM-5.1", "结论"],
    [
        ("Generic free actor", "64.4%", "71.9%", "信息完整，但目标身份约束仍隐式"),
        ("Generic + mode", "75.0%", "75.0%", "只增加保持/重选标签仍不足"),
        ("Generic + validity gate", "65.0%", "73.1%", "有效性检查不能阻止换到另一个有效对象"),
        ("Untyped pre-refresh plan", "81.2%", "70.6%", "普通提前规划不稳定"),
        ("Exact CTA", "95.0%", "96.2%", "刷新前编译目标承诺"),
        ("Lifecycle free actor", "96.9%", "98.1%", "类型化状态、无确定性 gate"),
        ("Lifecycle + gate", "98.1%", "100.0%", "gate 额外提升 1.2--1.9 点"),
    ], [2700, 1350, 1350, 3960], small=8.5)
base.add_body(doc, "Generic 已经保存关键事实，但两个模型仍有约三成任务失败；CTA 将准确率提高到 95% 以上。只增加 mode 字段、普通提前规划或动作有效性 gate 都不能稳定达到 CTA 的效果，因此当前证据把主要收益定位到刷新前目标承诺编译。")

doc.add_heading("对称策略控制", level=2)
base.add_body(doc, "Always-Lock + validity 与 Always-Reevaluate 在 160 条任务上都为 96/160（60%），但错误方向相反：前者在保持型任务为 80/80、重选型为 16/80；后者为 16/80 和 80/80。这说明问题不能用“始终保存旧 ID”或“始终使用最新结果”解决。")

# Replication, baselines, consequences
doc.add_heading("5. 独立复制、强基线与实际写入后果", level=1)
doc.add_heading("5.1 240 条三模型独立复制", level=2)
base.add_table(doc,
    ["模型", "Generic", "CTA", "Generic 条件漂移", "CTA 条件漂移", "差值"],
    [
        ("Qwen3.5", "47.5%", "70.8%", "43/72", "0", "+23.3 点"),
        ("GLM-5.1", "70.0%", "94.2%", "38/80", "0", "+24.2 点"),
        ("DeepSeek", "73.8%", "91.2%", "59/79", "0", "+17.5 点"),
    ], [1450, 1250, 1250, 1950, 1750, 1710], small=8.7)
base.add_body(doc, "三模型的 Generic 合计出现 140 次核心漂移，CTA 在相应严格机会中均为 0。逐领域和逐模板家族剔除后，CTA 相对 Generic 的差值始终为正，说明结果不是由单一领域或措辞模板主导。")

doc.add_heading("5.2 完整历史与刷新后强提醒", level=2)
base.add_table(doc,
    ["模型", "普通完整历史", "刷新后强提醒", "CTA"],
    [
        ("Qwen3.5", "63.3%", "69.6%", "70.8%"),
        ("GLM-5.1", "67.1%", "80.8%", "94.2%"),
        ("DeepSeek", "68.8%", "75.8%", "91.2%"),
    ], [1900, 2500, 2700, 2260], small=8.8)
base.add_body(doc, "完整历史和刷新后的强语义提醒都有一定帮助，但在每个模型的 80 条保持型 changed-winner 任务中，强提醒仍分别替换目标 56、38、42 次。CTA 在 GLM 和 DeepSeek 上明显更强，在 Qwen 上与强提醒总体持平。")

doc.add_heading("5.3 SQLite 实际后果", level=2)
base.add_table(doc,
    ["模型", "普通完整历史 wrong writes", "强提醒 wrong writes"],
    [
        ("Qwen3.5", "87/240", "70/240"),
        ("GLM-5.1", "79/240", "46/240"),
        ("DeepSeek", "75/240", "57/240"),
        ("合计", "241", "173；两类共 414 次"),
    ], [2000, 3600, 3760], small=8.7)
base.add_body(doc, "六组完整历史输出共 1,440 个 episode，错误会修改 SQLite 中不应被修改的记录。v7 中 Generic 的 43、38、59 次核心漂移也全部成为 wrong-entity writes。CTA 仍可能因初始 ID、mode 或 selector grounding 错误产生非核心错误，因此 0 条件漂移不等于 0 工具风险。")

# Humans and external
doc.add_heading("6. 人工验证与外部环境结果", level=1)
doc.add_heading("6.1 人工语义验证", level=2)
base.add_table(doc,
    ["验证", "规模", "结果"],
    [
        ("盲标一致性", "3 人 × 100 条", "Fleiss kappa=.708；Krippendorff alpha=.709"),
        ("多数票与 gold", "100 条", "86%；94 条确定多数中为 91.5%"),
        ("保持/重选核心", "可执行保持 + 动态", "保持可执行 86.7%；动态 98%"),
        ("独立英文改写", "50 条 × 2 模型", "Generic 60/74；CTA 90/98"),
    ], [2300, 2300, 4760], small=8.7)
base.add_body(doc, "人类结果支持 Preserve/Reevaluate 是可操作的语义区别，且模型差距在独立改写上继续存在。失效目标后的 Reject 策略一致性较低，因此该部分在论文中作为规范性 fallback，而不是核心人类语义结论。")

doc.add_heading("6.2 公开 benchmark 与工具闭环", level=2)
base.add_table(doc,
    ["对象", "范围", "结果"],
    [
        ("ToolSandbox", "129 个语义家族", "严格原生机会 0；1 个 near-match"),
        ("AppWorld", "732 任务 / 244 家族", "严格机会 0；公开轨迹未见后绑定替换"),
        ("tau3", "2,449 任务 / 10,832 轨迹", "严格原生机会 0"),
        ("ToolSandbox-compatible", "24 条自定义干预", "GLM Generic 3/6；Stable 0/2"),
        ("低干预外部闭环", "96 条扩展 + AppWorld", "正确绑定后的条件 TRI 为 0"),
    ], [2300, 3000, 4060], small=8.5)
base.add_body(doc, "公开任务通常缺少正确绑定后的独立状态变化、同角色竞争赢家或可评分的错误写入，因此现有审计主要说明 benchmark 覆盖盲区。24 条 ToolSandbox-compatible 任务表明，在真实数据库和工具接口中人为构造严格机会时现象可以发生；但自然度更高、干预更低的外部闭环为零结果。")
base.add_callout(doc, "外部结论", "目前外部证据支持“机会覆盖有限、现象依赖模型与控制器”，不支持估计真实部署中的发生率。", base.LIGHT_ORANGE, base.ORANGE)

# Paper status
doc.add_heading("7. 当前论文进展", level=1)
base.add_body(doc, "当前主线已经固定为：状态刷新更新世界知识，但不自动授权 Agent 改变已经绑定的操作目标；普通结构化状态与完整历史仍可能发生 post-binding drift；刷新前目标承诺编译能够在受控单目标任务中显著减少该错误。")
doc.add_heading("7.1 论文和实验材料现状", level=2)
base.add_table(doc,
    ["材料", "当前状态"],
    [
        ("论文材料", "主文已形成 7 页正文 + 1 页参考文献版本；补充材料已整理方法接口、统计口径和额外审计"),
        ("实验与溯源", "v3、v7 三模型、强基线、SQLite、人类及外部结果已进入证据链，可回到冻结数据和 raw runs"),
        ("当前收尾", "继续核对主文数字、统一术语、检查匿名材料并冻结投稿包"),
    ], [2100, 7260], small=8.8)

doc.add_heading("7.2 当前结论边界", level=2)
for item in [
    "结果支持受控、单刷新、单操作目标场景中的 TRI 机制，不代表所有 Agent 都会出现该问题；",
    "CTA 在 GLM 和 DeepSeek 上明显优于刷新后强提醒，在 Qwen 上总体持平；",
    "多刷新、多角色尚未跨模型稳定解决；TRI 与 Binding Drift 的区别在于先条件化到初始绑定正确，再研究状态变化后是否得到用户授权重选。",
]:
    base.add_bullet(doc, item)

doc.core_properties.title = "TRI 论文研究进展同步"
doc.core_properties.subject = "时序指称完整性论文的研究思路、实验结果与当前进展"
doc.core_properties.author = "TRI Project"
doc.core_properties.keywords = "TRI, Temporal Referent Integrity, Tool-Using Agents, AAAI 2027"

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
