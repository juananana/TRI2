#!/usr/bin/env python3
"""Build the formal Chinese TRI research summary as a polished DOCX."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "planning" / "TRI_formal_research_summary_zh.md"
OUTPUT = ROOT / "output" / "TRI论文研究总结_20260721.docx"

FONT_LATIN = "Arial"
FONT_CJK = "Hiragino Sans GB"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 42, 54)
MUTED = RGBColor(95, 105, 115)
LIGHT_FILL = "F2F4F7"
GRID = "C9D1DA"
CONTENT_DXA = 9360


def set_run_font(run, size=None, bold=None, italic=None, color=None, cjk=FONT_CJK):
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), cjk)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), cjk)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hint"), "eastAsia")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, size, color=INK, bold=False):
    style.font.name = FONT_LATIN
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CJK)
    style._element.get_or_add_rPr().rFonts.set(qn("w:cs"), FONT_CJK)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hint"), "eastAsia")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold


def configure_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, 10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (11.5, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        set_style_font(style, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.08

    for name in ("List Number", "List Bullet"):
        style = doc.styles[name]
        set_style_font(style, 10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.18


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def configure_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    margins = {"top": 90, "bottom": 90, "start": 120, "end": 120}
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in margins.items():
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), GRID)


def column_widths(count):
    if count == 2:
        return [4300, 5060]
    if count == 3:
        return [4200, 2580, 2580]
    if count == 4:
        return [2900, 2100, 2100, 2260]
    return [CONTENT_DXA // count] * count


def add_inline(paragraph, text, default_size=10.5, default_color=INK):
    pattern = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, default_size, color=default_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, default_size, bold=True, color=default_color)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, default_size, italic=True, color=default_color)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, default_size - 0.3, color=DARK_BLUE, cjk="Menlo")
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, default_size, color=default_color)


def add_table(doc, rows):
    count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=count)
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for r_index, row in enumerate(rows):
        for c_index in range(count):
            value = row[c_index] if c_index < len(row) else ""
            cell = table.cell(r_index, c_index)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, value, 9.2, INK)
            if r_index == 0:
                set_cell_shading(cell, LIGHT_FILL)
                for run in p.runs:
                    run.bold = True
    configure_table_geometry(table, column_widths(count))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, end])
    set_run_font(run, 9, color=MUTED)


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("TRI 论文研究总结")
    set_run_font(run, 9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    add_page_number(footer)


def add_masthead(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("研究总结")
    set_run_font(run, 10, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("更新世界状态不等于重新绑定操作目标")
    set_run_font(run, 23, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(15)
    run = p.add_run("工具型智能体中的时序指称完整性")
    set_run_font(run, 14, color=DARK_BLUE)

    metadata = [
        ("英文题目", "Updating the World Is Not Rebinding the Target: Temporal Referent Integrity in Tool-Using Agents"),
        ("投稿方向", "AAAI 2027 Main Technical Track"),
        ("更新时间", "2026 年 7 月 21 日"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{label}：")
        set_run_font(run, 10, bold=True, color=MUTED)
        run = p.add_run(value)
        set_run_font(run, 10, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(14)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "2E74B5")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def parse_markdown(doc):
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = lines.index("## 摘要")
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            index += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            index += 1
            continue
        if line.startswith("| ") and index + 1 < len(lines) and lines[index + 1].strip().startswith("|---"):
            rows = []
            rows.append([cell.strip() for cell in line.strip("|").split("|")])
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
                index += 1
            add_table(doc, rows)
            continue
        if re.match(r"^\d+\. ", line):
            while index < len(lines) and re.match(r"^\d+\. ", lines[index].strip()):
                item = re.sub(r"^\d+\. ", "", lines[index].strip())
                p = doc.add_paragraph(style="List Number")
                add_inline(p, item)
                index += 1
            continue
        paragraph = [line]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate or candidate.startswith("#") or candidate.startswith("|") or re.match(r"^\d+\. ", candidate):
                break
            paragraph.append(candidate)
            index += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(paragraph))


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    add_masthead(doc)
    parse_markdown(doc)
    doc.core_properties.title = "TRI 论文研究总结"
    doc.core_properties.subject = "工具型智能体中的时序指称完整性"
    doc.core_properties.author = "TRI Research Project"
    doc.core_properties.keywords = "Temporal Referent Integrity; Tool-Using Agents; AAAI 2027"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
